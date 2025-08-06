import os
import io
import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI

# --- Load environment variables ---
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
app_password = os.getenv("STREAMLIT_PASSWORD")
client = OpenAI(api_key=api_key)

# --- Helper Functions ---
def extract_text_from_pdf(file):
    text = ""
    pdf = fitz.open(stream=file.read(), filetype="pdf")
    for page in pdf:
        text += page.get_text()
    return text

def summarize_text(text):
    prompt = f"""
    Summarize the following legal document with detailed sections.
    If any section is missing, say "Not specified".

    Sections to include:
    1. Parties
    2. Effective Date
    3. Term
    4. Confidential Information
    5. Obligations
    6. Jurisdiction
    7. Risk Flags

    Document:
    {text[:10000]}
    """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

def download_summary(summary_text):
    doc = Document()
    doc.add_heading("Legal Document Summary", level=1)
    doc.add_paragraph(summary_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(
        label="📥 Download Summary (.docx)",
        data=buffer,
        file_name="summary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- Password Protection ---
st.write("🔐 STREAMLIT_PASSWORD loaded:", bool(app_password))
entered_password = st.text_input("🔒 Enter Access Password:", type="password")
if entered_password != app_password:
    st.warning("Please enter the correct password.")
    st.stop()

# --- UI Title and Mode Selection ---
st.title("Legal PDF Summarizer")
mode = st.radio("What would you like to do?", [
    "🟢 🔍 Summarize a Single Document",
    "🟢 📊 Summarize and Compare Multiple Documents"
])

# --- Upload & Trigger Logic ---
if mode.startswith("🟢 🔍"):
    uploaded_file = st.file_uploader("Upload a PDF", type="pdf", key="single")

    if st.button("🚀 Start"):
        if uploaded_file:
            st.info("Processing single document...")
            pdf_text = extract_text_from_pdf(uploaded_file)
            summary = summarize_text(pdf_text)
            st.subheader("Summary")
            st.write(summary)
            download_summary(summary)
        else:
            st.warning("Please upload a PDF file.")

else:
    uploaded_files = st.file_uploader("Upload 2–5 PDFs", type="pdf", accept_multiple_files=True, key="multi")

    if st.button("🚀 Start"):
        if uploaded_files and 2 <= len(uploaded_files) <= 5:
            st.info("Processing multiple documents...")
            summaries = []
            for file in uploaded_files:
                pdf_text = extract_text_from_pdf(file)
                summary = summarize_text(pdf_text)
                summaries.append((file.name, summary))

            for name, summary in summaries:
                st.subheader(f"📄 {name}")
                st.write(summary)
            # Optionally: add merged .docx export here
        else:
            st.warning("Please upload between 2 and 5 PDF files.")
